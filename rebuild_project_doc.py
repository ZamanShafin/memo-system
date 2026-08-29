import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def apply_doc_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

def add_callout_box(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.8)
    set_cell_background(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if title:
        r_title = p.add_run(f"{title}\n")
        r_title.bold = True
        r_title.font.size = Pt(10)
        r_title.font.color.rgb = RGBColor(30, 27, 75)
    r_body = p.add_run(text)
    r_body.font.size = Pt(9.5)
    r_body.font.color.rgb = RGBColor(51, 65, 85)
    doc.add_paragraph() # Spacing

def add_diagram_box(doc, diagram_text, caption=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.8)
    set_cell_background(cell, "0F172A") # Slate 900 dark background
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(diagram_text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(226, 232, 240) # Slate 200 light text
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run(f"Figure: {caption}")
        r_cap.font.size = Pt(9)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

# Create Document
doc = docx.Document()
apply_doc_margins(doc)

# -----------------------------------------------------------------------------
# TITLE & HEADER
# -----------------------------------------------------------------------------
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_main = p_title.add_run("CSE226 Final Project Technical Documentation\n")
r_main.bold = True
r_main.font.size = Pt(20)
r_main.font.color.rgb = RGBColor(30, 27, 75)

r_sub = p_title.add_run("Multi-Tenant Inter-Office Memo Management System\nDeterministic Sequential Approval Workflow Engine & Document Governance SaaS")
r_sub.font.size = Pt(11.5)
r_sub.font.bold = True
r_sub.font.color.rgb = RGBColor(79, 70, 229)

# METADATA TABLE
meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Author / Student:", "Zaman Shafin (North South University)"),
    ("Course & Term:", "CSE226 Fundamentals of Vibe Coding — Summer 2026"),
    ("Deployed System URL:", "https://memo-system-pjbj.vercel.app"),
    ("Source Code Archive:", "https://github.com/ZamanShafin/memo-system/raw/main/source_code.zip"),
    ("GitHub Repository:", "https://github.com/ZamanShafin/memo-system")
]
for i, (k, v) in enumerate(meta_data):
    c1, c2 = meta_table.rows[i].cells
    c1.width = Inches(2.0)
    c2.width = Inches(4.8)
    set_cell_background(c1, "F1F5F9")
    c1.paragraphs[0].add_run(k).bold = True
    c1.paragraphs[0].runs[0].font.size = Pt(9.5)
    c2.paragraphs[0].add_run(v).font.size = Pt(9.5)

doc.add_paragraph()

# -----------------------------------------------------------------------------
# SECTION 1: SYSTEM OVERVIEW
# -----------------------------------------------------------------------------
h1 = doc.add_heading("1. System Overview & Core Capabilities", level=1)
h1.style.font.color.rgb = RGBColor(30, 27, 75)

doc.add_paragraph(
    "The Inter-Office Memo Management System is a full-stack, multi-tenant SaaS application built to digitize, "
    "govern, and audit corporate memorandum workflows. It eliminates fragmented email threads and physical paper slips "
    "by enforcing deterministic sequential approvals, active delegation, immutable version history, and cryptographic PDF generation."
)

add_callout_box(doc, 
    "• Multi-Tenant Isolation: Logical data partitioning with non-nullable org_id scoping across every database query.\n"
    "• Strict Turn Governance: Memos execute sequentially (Requester -> Dept Head -> Finance -> Director -> CEO). Later participants are physically blocked from acting early.\n"
    "• Active Delegation: Temporary date-bounded approval authority handoffs when officers are on leave.\n"
    "• Dynamic Reviewer Injection: Approvers can dynamically insert specialized reviewers (e.g. Legal Counsel) mid-stream.\n"
    "• Tamper-Evident Auditing: Append-only audit logs with timestamps, actor IDs, comments, and client IPs.\n"
    "• PDF Seal Engine: Generates official corporate PDFs complete with digital signatures and QR verification codes.",
    "Key System Highlights"
)

# -----------------------------------------------------------------------------
# SECTION 2: REQUIREMENTS IMPLEMENTED (COMPLIANCE MATRIX)
# -----------------------------------------------------------------------------
h2 = doc.add_heading("2. Requirements Implemented (Compliance Matrix)", level=1)
h2.style.font.color.rgb = RGBColor(30, 27, 75)

req_rows = [
    ("Req 1-4", "Multi-Tenancy & RBAC", "Complete org_id data perimeters, Admin/User roles, department management."),
    ("Req 5-7", "Auth & Memo Authoring", "JWT tokens, password reset, Quill.js rich text editor, priority tags, file uploads."),
    ("Req 8-10", "Sequential Workflow Engine", "Strict linear turn execution, workflow templates, mid-stream reviewer insertion."),
    ("Req 11-13", "Decisions & Delegation", "Approve, Reject, Request Changes with audit notes; date-bounded delegation."),
    ("Req 14-17", "PDFs, Alerts & Versions", "Cryptographic PDF seal generation, notification tray, immutable version snapshots."),
    ("Req 18-22", "Audit & Security", "Append-only audit logs, turnaround charts, admin self-deactivation protection."),
    ("Req 23-30", "Deployment & Testing", "Live on Vercel, sanitized source ZIP, 20/20 passing pytest suite, full prompt log.")
]

t_req = doc.add_table(rows=len(req_rows)+1, cols=3)
t_req.alignment = WD_TABLE_ALIGNMENT.CENTER
r_hdrs = ["Ref #", "Requirement Area", "Implementation Highlights"]
for j, h in enumerate(r_hdrs):
    c = t_req.rows[0].cells[j]
    set_cell_background(c, "1E1B4B")
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(255, 255, 255)

for i, (r_id, r_area, r_desc) in enumerate(req_rows):
    row = t_req.rows[i+1]
    for col_idx, txt in enumerate([r_id, r_area, r_desc]):
        c = row.cells[col_idx]
        if i % 2 == 0:
            set_cell_background(c, "F8FAFC")
        c.paragraphs[0].add_run(txt).font.size = Pt(9)

doc.add_paragraph()

# -----------------------------------------------------------------------------
# SECTION 3: TECHNOLOGY STACK
# -----------------------------------------------------------------------------
h3 = doc.add_heading("3. Technology Stack", level=1)
h3.style.font.color.rgb = RGBColor(30, 27, 75)

stack_rows = [
    ("Backend API Framework", "FastAPI (Asynchronous Python 3.10+ REST framework)"),
    ("Database Engine & ORM", "PostgreSQL 16 (Neon Serverless Cloud DB) + SQLAlchemy 2.0 ORM"),
    ("Authentication & Security", "PyJWT (JSON Web Tokens) + Passlib (PBKDF2-SHA256 & Bcrypt hashing)"),
    ("Frontend Single Page App", "Vanilla JavaScript SPA engine (app.js) with zero build-step overhead"),
    ("Styling & Design System", "Tailwind CSS utility framework + Lucide Icons + Google Fonts"),
    ("Rich Text Authoring", "Quill.js WYSIWYG editor with custom toolbar & HTML sanitization"),
    ("Visualizations & Charts", "Chart.js for real-time status doughnuts and turnaround duration bars"),
    ("Document / PDF Engine", "ReportLab & xhtml2pdf with embedded signature stamps and QR codes"),
    ("Cloud Infrastructure", "Vercel Serverless Functions with automated CI/CD pipeline"),
    ("Testing & Tooling", "Pytest automated test harness (20/20 tests passing) + Astral uv")
]

t_stack = doc.add_table(rows=len(stack_rows)+1, cols=2)
t_stack.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Layer / Component", "Technology & Libraries"]):
    c = t_stack.rows[0].cells[j]
    set_cell_background(c, "312E81")
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(255, 255, 255)

for i, (layer, tech) in enumerate(stack_rows):
    row = t_stack.rows[i+1]
    for col_idx, txt in enumerate([layer, tech]):
        c = row.cells[col_idx]
        if i % 2 == 0:
            set_cell_background(c, "F8FAFC")
        c.paragraphs[0].add_run(txt).font.size = Pt(9)

doc.add_page_break()

# -----------------------------------------------------------------------------
# SECTION 4: SYSTEM ARCHITECTURE & DIAGRAM
# -----------------------------------------------------------------------------
h4 = doc.add_heading("4. System Architecture & Diagram", level=1)
h4.style.font.color.rgb = RGBColor(30, 27, 75)

doc.add_paragraph(
    "The application is structured into a clean 3-tier decoupled architecture comprising a Client SPA, "
    "a Stateless FastAPI REST API service layer, and a cloud-native PostgreSQL relational database."
)

arch_diagram = """+-----------------------------------------------------------------------------+
|                            CLIENT TIER (Browser / Mobile)                   |
|  - Vanilla JS SPA Router (showView)           - Tailwind Responsive Layout  |
|  - Reactive State Store (appState)            - Quill.js WYSIWYG Editor     |
|  - Interactive Chart.js Visualizations        - Lucide Icon System          |
+--------------------------------------+--------------------------------------+
                                       | HTTP REST / Bearer JWT
                                       v
+-----------------------------------------------------------------------------+
|                     APPLICATION & REST API TIER (FastAPI)                   |
|  +-------------------+  +--------------------+  +------------------------+  |
|  |   /auth Router    |  |   /memos Router    |  |   /workflow Router     |  |
|  | - JWT Auth & Reset|  | - CRUD & Attachments| | - Action & Turn Engine |  |
|  +-------------------+  +--------------------+  +------------------------+  |
|  +-------------------+  +--------------------+  +------------------------+  |
|  |   /admin Router   |  |   /reports Router  |  |  /delegations Router   |  |
|  | - Users & Depts   |  | - Analytics Stats  |  | - Date-Bounded Rules   |  |
|  +-------------------+  +--------------------+  +------------------------+  |
+--------------------------------------+--------------------------------------+
                                       | Internal Service Invocations
                                       v
+-----------------------------------------------------------------------------+
|                        CORE BUSINESS LOGIC SERVICES                         |
|  - Workflow State Machine        - Immutable Version Snapshotting           |
|  - Cryptographic PDF & QR Engine - Real-Time In-App Notification Dispatcher |
|  - Append-Only Audit Logger      - Role-Based Access Control (RBAC) Guard   |
+--------------------------------------+--------------------------------------+
                                       | SQLAlchemy 2.0 (SSL / Connection Pool)
                                       v
+-----------------------------------------------------------------------------+
|                     DATA & STORAGE TIER (Neon PostgreSQL)                   |
|  [organizations]  [users]        [departments]  [memo_categories]           |
|  [memos]          [memo_steps]   [memo_versions][workflow_delegations]      |
|  [memo_comments]  [attachments]  [audit_logs]   [notifications]             |
+-----------------------------------------------------------------------------+"""

add_diagram_box(doc, arch_diagram, "High-Level 3-Tier System Architecture Diagram")

# -----------------------------------------------------------------------------
# SECTION 5: DATABASE DESIGN & MULTI-TENANCY ENFORCEMENT
# -----------------------------------------------------------------------------
h5 = doc.add_heading("5. Database Design & Multi-Tenancy Enforcement", level=1)
h5.style.font.color.rgb = RGBColor(30, 27, 75)

doc.add_paragraph(
    "Multi-tenancy is enforced through strict logical data partitioning. Every single table in the database "
    "contains a non-nullable foreign key referencing organizations.id. Multi-tenancy is guaranteed at the database "
    "and API layer through the following mechanisms:"
)

doc.add_paragraph("• Server-Side Query Scoping: Every database query in all routers automatically includes models.Entity.org_id == current_user.org_id. Client-side filtering is never relied upon for security.", style="List Bullet")
doc.add_paragraph("• Direct Object Access Validation: When accessing a resource by ID (e.g. GET /memos/10), the server verifies memo.org_id == current_user.org_id. If the tenant IDs mismatch, an immediate HTTP 404/403 is thrown.", style="List Bullet")
doc.add_paragraph("• Cascade Isolation: Deleting or modifying a department, user, or workflow step is strictly constrained to the user's organization.", style="List Bullet")

db_erd = """+-------------------+        +-------------------+        +----------------------+
|   ORGANIZATION    | 1    * |       USER        | 1    * |         MEMO         |
|-------------------|--------|-------------------|--------|----------------------|
| id (PK)           |        | id (PK)           |        | id (PK)              |
| name, code        |        | org_id (FK)       |        | org_id (FK)          |
| contact_email     |        | department_id(FK) |        | memo_number, title   |
+-------------------+        | full_name, email  |        | body_html, priority  |
                             | role, is_active   |        | status, author_id    |
                             +-------------------+        | current_assignee_id  |
                                                          +----------+-----------+
                                                                     | 1
                                                                     | *
+-----------------------+    +-----------------------+    +----------v-----------+
|    MEMO_VERSION       |    |   WORKFLOW_DELEGATION |    |  MEMO_WORKFLOW_STEP  |
|-----------------------|    |-----------------------|    |----------------------|
| id (PK), memo_id (FK) |    | id (PK), org_id (FK)  |    | id (PK), memo_id(FK) |
| version_num, body_html|    | delegator_id (FK)     |    | step_index, role_name|
| change_summary        |    | delegatee_id (FK)     |    | assigned_user_id (FK)|
| created_by, created_at|    | start_date, end_date  |    | status, decision     |
+-----------------------+    +-----------------------+    +----------------------+"""

add_diagram_box(doc, db_erd, "Entity-Relationship & Multi-Tenancy Diagram")

doc.add_page_break()

# -----------------------------------------------------------------------------
# SECTION 6: WORKFLOW DESIGN & STATE TRANSITIONS
# -----------------------------------------------------------------------------
h6 = doc.add_heading("6. Sequential Workflow Design & Turn Governance", level=1)
h6.style.font.color.rgb = RGBColor(30, 27, 75)

doc.add_paragraph(
    "The sequential workflow engine operates as a deterministic finite-state machine with strict turn-based locks. "
    "Only the currently assigned participant (or their authorized active delegate) can sign off."
)

wf_diagram = """    [ Author Creates Memo ] 
               |
               v
         [ DRAFT STATE ]  <====================================+
               |                                               |
        Author Submits                                         |
               v                                               |
     +-------------------------------------------------------+ |
     |           PENDING APPROVAL / REVIEW STATE             | |
     |                                                       | |
     |  Step 1: Dept Head  ===> Approve ===> Step 2          | |
     |  Step 2: Finance    ===> Approve ===> Step 3          | |
     |  Step 3: Director   ===> Approve ===> Step 4          | |
     |  Step 4: CEO        ===> Approve ===> [ APPROVED ]    | |
     +---------+--------------------+------------------------+ |
               |                    |                          |
        Reviewer Rejects      Request Changes                  |
               |                    |                          |
               v                    +==========================+
         [ REJECTED ]              (Returns to Author for edits;
     (Workflow Terminated)          Creates immutable MemoVersion snapshot)"""

add_diagram_box(doc, wf_diagram, "Sequential Approval Workflow State Machine Diagram")

# Workflow Rules Table
wf_rules = [
    ("Draft -> Submission", "Author drafts memo and defines sequence. Submitting advances turn to Step 1 and dispatches notifications."),
    ("Step-by-Step Turns", "Step N must be completed before Step N+1 can act. Later approvers are physically blocked by server-side checks."),
    ("Dynamic Reviewer Insertion", "Approver can select 'Approve & Add Reviewer' to insert a specialist (e.g. Legal Counsel) before next step."),
    ("Decline & Reroute", "Turn holder can reassign their active step to an eligible colleague in the same organization."),
    ("Delegation Protocol", "When active date window is valid, designated delegatee can sign off on behalf of delegator with full audit attribution."),
    ("Changes Requested Cycle", "Reviewer returns memo with comments. Author updates content; system creates MemoVersion snapshot before resubmission."),
    ("Final Approval & Seal", "Final approver (CEO) marks workflow Approved. System locks content as read-only and generates official signed PDF.")
]

t_wf = doc.add_table(rows=len(wf_rules)+1, cols=2)
t_wf.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Workflow Stage / Rule", "Operational Behavior & Governance"]):
    c = t_wf.rows[0].cells[j]
    set_cell_background(c, "1E1B4B")
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(255, 255, 255)

for i, (stage, desc) in enumerate(wf_rules):
    row = t_wf.rows[i+1]
    for col_idx, txt in enumerate([stage, desc]):
        c = row.cells[col_idx]
        if i % 2 == 0:
            set_cell_background(c, "F8FAFC")
        c.paragraphs[0].add_run(txt).font.size = Pt(9)

doc.add_paragraph()

# -----------------------------------------------------------------------------
# SECTION 7: SECURITY ARCHITECTURE
# -----------------------------------------------------------------------------
h7 = doc.add_heading("7. Security Architecture & Access Control", level=1)
h7.style.font.color.rgb = RGBColor(30, 27, 75)

doc.add_paragraph("• Cryptographic Passwords: Passwords are salted and encrypted using PBKDF2-SHA256 and Bcrypt algorithms.\n"
                  "• JWT Bearer Sessions: Stateless tokens signed with HMAC-SHA256 with 24-hour expiration.\n"
                  "• Admin Self-Protection: Dedicated server-side rules block administrators from deactivating or demoting their own accounts.\n"
                  "• Injection & XSS Defenses: Parameterized queries via SQLAlchemy prevent SQL injection; rich text HTML is sanitized before rendering.\n"
                  "• Secure Cloud Connection: Database connection strictly requires SSL/TLS encryption (sslmode=require).")

# -----------------------------------------------------------------------------
# SECTION 8: KNOWN LIMITATIONS
# -----------------------------------------------------------------------------
h8 = doc.add_heading("8. Known Limitations (Documented Transparently)", level=1)
h8.style.font.color.rgb = RGBColor(30, 27, 75)

doc.add_paragraph("1. In-App Notification Transport: Uses responsive client-side HTTP polling rather than full-duplex WebSockets.\n"
                  "2. Object Storage: File attachments are persisted within the PostgreSQL database layer rather than external AWS S3 buckets.\n"
                  "3. Two-Factor Authentication: MFA/TOTP is reserved for future enterprise roadmap versions.")

# -----------------------------------------------------------------------------
# SECTION 9: DEPLOYMENT INFORMATION & DEMO CREDENTIALS
# -----------------------------------------------------------------------------
h9 = doc.add_heading("9. Deployment Information & Demonstration Accounts", level=1)
h9.style.font.color.rgb = RGBColor(30, 27, 75)

doc.add_paragraph("• Live Deployed URL: https://memo-system-pjbj.vercel.app\n"
                  "• Source Code ZIP: https://github.com/ZamanShafin/memo-system/raw/main/source_code.zip\n"
                  "• Active Organization Code: acme\n"
                  "• Universal Password: password123")

demo_accounts = [
    ("Admin", "admin@acmecorp.com", "Sarah Jenkins (System Administrator)"),
    ("Author", "alex.morgan@acmecorp.com", "Alex Morgan (Senior Engineer / Author)"),
    ("Dept Head", "head.eng@acmecorp.com", "David Vance (VP of Engineering / Approver)"),
    ("Delegate", "jessica.taylor@acmecorp.com", "Jessica Taylor (Acting Delegate for David)"),
    ("Finance", "finance.mgr@acmecorp.com", "Rachel Green (Chief Financial Manager)"),
    ("Director", "director@acmecorp.com", "Marcus Sterling (Director of Operations)"),
    ("CEO", "ceo@acmecorp.com", "Eleanor Vance (Chief Executive Officer)")
]

t_demo = doc.add_table(rows=len(demo_accounts)+1, cols=3)
t_demo.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Role", "Login Email", "Persona / Purpose"]):
    c = t_demo.rows[0].cells[j]
    set_cell_background(c, "312E81")
    r = c.paragraphs[0].add_run(h)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(255, 255, 255)

for i, (r_role, r_em, r_purp) in enumerate(demo_accounts):
    row = t_demo.rows[i+1]
    for col_idx, txt in enumerate([r_role, r_em, r_purp]):
        c = row.cells[col_idx]
        if i % 2 == 0:
            set_cell_background(c, "F8FAFC")
        c.paragraphs[0].add_run(txt).font.size = Pt(9)

# Save Document
doc.save("final_submission/PROJECT_DOCUMENTATION.docx")
doc.save("submission/PROJECT_DOCUMENTATION.docx")
doc.save("PROJECT_DOCUMENTATION.docx")

print("SUCCESS: Rebuilt PROJECT_DOCUMENTATION.docx with clear brief format, Architecture Diagram, and Workflow State Diagram!")
