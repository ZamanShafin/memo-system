import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def generate_prd():
    doc = docx.Document()
    
    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def set_cell_background(cell, fill_color):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
        tcPr.append(shd)

    def add_section_header(title, level=1):
        h = doc.add_heading(title, level=level)
        h.style.font.color.rgb = RGBColor(30, 27, 75)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        return h

    # -------------------------------------------------------------
    # 1. TITLE & COVER HEADER
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('PRODUCT REQUIREMENTS DOCUMENT (PRD)\n')
    r_title.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = RGBColor(30, 27, 75) # Indigo 950

    r_sub = p_title.add_run('Enterprise Multi-Tenant Inter-Office Memo Management System\nDeterministic Sequential Approval & Document Governance SaaS Platform')
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600

    doc.add_paragraph()

    # DOCUMENT CONTROL TABLE
    ctrl_table = doc.add_table(rows=7, cols=2)
    ctrl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ctrl_data = [
        ('Project Title:', 'Multi-Tenant Inter-Office Memo Management System'),
        ('Product Version:', 'v4.3.0 Enterprise Release (Production Live)'),
        ('Document Status:', 'Approved & Baseline Implemented'),
        ('Product Owner / Lead:', 'Zaman Shafin (North South University)'),
        ('Course / Context:', 'CSE226 Foundations of Vibe Coding'),
        ('Live Cloud Deployment:', 'https://memo-system-pjbj.vercel.app'),
        ('Source Code Repository:', 'https://github.com/ZamanShafin/memo-system')
    ]
    for i, (k, v) in enumerate(ctrl_data):
        cell_k, cell_v = ctrl_table.rows[i].cells
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, 'F1F5F9')
        p_k = cell_k.paragraphs[0]
        r_k = p_k.add_run(k)
        r_k.bold = True
        r_k.font.size = Pt(9.5)
        p_v = cell_v.paragraphs[0]
        r_v = p_v.add_run(v)
        r_v.font.size = Pt(9.5)

    doc.add_page_break()

    # -------------------------------------------------------------
    # 2. EXECUTIVE SUMMARY & BUSINESS OBJECTIVES
    # -------------------------------------------------------------
    add_section_header('1. Executive Summary & Business Objectives', level=1)
    
    doc.add_paragraph(
        '1.1 Problem Statement:\n'
        'In traditional modern enterprise environments, internal memorandum approvals, resource requisitions, policy updates, and executive sign-offs rely heavily on informal communication channels—primarily email chains, physical paperwork, and instant messaging apps. This produces significant organizational liabilities: lack of strict turn-based governance, lost audit trails, unauthorized approvals, bottleneck opacity, and zero version history during revisions.'
    )
    
    doc.add_paragraph(
        '1.2 Product Vision & Solution:\n'
        'The Inter-Office Memo Management System is a centralized, cloud-native Software-as-a-Service (SaaS) platform providing strict sequential workflow execution, logical multi-tenant data isolation, dynamic reviewer injection, date-bounded delegation, tamper-evident audit logging, and cryptographic PDF seal generation.'
    )

    doc.add_paragraph('1.3 Business Objectives & Success Metrics (KPIs):')
    kpis = [
        ('Turnaround Efficiency', 'Reduce average inter-office memorandum approval time from 7+ business days to under 24 hours via automated turn handoffs and in-app alerts.'),
        ('Governance & Compliance', 'Ensure 100% auditable chain-of-custody for all financial, operational, and administrative decisions with immutable timestamped logs.'),
        ('Tenant Security & Isolation', 'Zero data leakage across organizations sharing the same cloud database infrastructure via universal tenant-scoped data perimeters.'),
        ('Mobile & Touch Accessibility', 'Enable executive officers to review, comment, and sign off memoranda from any mobile, tablet, or desktop device with responsive layouts.')
    ]
    for kpi, desc in kpis:
        p = doc.add_paragraph(style='List Bullet')
        r_b = p.add_run(f'{kpi}: ')
        r_b.bold = True
        p.add_run(desc)

    # -------------------------------------------------------------
    # 3. USER PERSONAS & ROLE MATRIX
    # -------------------------------------------------------------
    add_section_header('2. User Personas & Role Matrix', level=1)
    
    doc.add_paragraph('The platform models an authentic corporate organizational hierarchy:')
    
    persona_table = doc.add_table(rows=8, cols=4)
    persona_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Persona / Name', 'Role & Dept', 'Primary Goals', 'System Permissions']
    for j, h in enumerate(headers):
        cell = persona_table.rows[0].cells[j]
        set_cell_background(cell, '1E1B4B')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    personas = [
        ('Sarah Jenkins', 'System Admin (Exec Ops)', 'Tenant governance, staff management, department config, audit inspection.', 'Full administrative access across organization, department deletion, user provisioning.'),
        ('Alex Morgan', 'Team Lead (Engineering)', 'Drafting memoranda, attaching technical files, tracking real-time approval status.', 'Create/edit drafts, submit workflows, view own memos, resubmit change requests.'),
        ('David Vance', 'Dept Head (Engineering)', 'Reviewing technical proposals, requesting revisions, assigning temporary delegates.', 'Approve/Reject Tier-1 steps, request changes, add intermediate reviewers, delegate authority.'),
        ('Jessica Taylor', 'Operations Specialist', 'Acting on behalf of Dept Head during leave with full delegated authority.', 'Execute approval actions on behalf of delegator when active date window is valid.'),
        ('Rachel Green', 'Finance Manager', 'Vetting expenditure, reviewing procurement budgets, preventing fiscal overrun.', 'Tier-2 budget sign-off, financial comments, change requests, rejection authority.'),
        ('Marcus Sterling', 'Operations Director', 'Operational oversight, multi-department coordination, strategic endorsements.', 'Tier-3 operational approvals, dynamic reviewer rerouting, executive notes.'),
        ('Eleanor Vance', 'CEO (Executive Office)', 'Final corporate authorization, formal sign-off, generating signed PDFs.', 'Final step approval, executive stamp, seals workflow into finalized archive.')
    ]

    for i, (p_name, p_role, p_goal, p_perm) in enumerate(personas):
        row = persona_table.rows[i+1]
        for col_idx, text in enumerate([p_name, p_role, p_goal, p_perm]):
            cell = row.cells[col_idx]
            if i % 2 == 0:
                set_cell_background(cell, 'F8FAFC')
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)

    doc.add_page_break()

    # -------------------------------------------------------------
    # 4. FUNCTIONAL REQUIREMENTS (FRs)
    # -------------------------------------------------------------
    add_section_header('3. Functional Requirements (FRs)', level=1)

    frs = [
        ('FR-1: Multi-Tenant Workspace & Org Registration',
         'The system shall allow new organizations to self-register with a unique tenant code. Every database record (users, memos, departments, categories, templates, logs) must store a non-nullable org_id ensuring complete data perimeter isolation.'),
        
        ('FR-2: User Identity & Role-Based Access Control (RBAC)',
         'The system shall support User and Admin roles. Admins manage staff, departments, and templates. The system shall enforce self-protection guards preventing admins from deactivating or demoting themselves.'),
        
        ('FR-3: Memorandum Creation & Rich-Text WYSIWYG',
         'Users shall create memos with automated sequential reference codes (MEMO-{ORG}-{YEAR}-{SEQ}), subject lines, categories, priority tags (Urgent, High, Normal), multi-file attachments, and rich formatted text via Quill.js.'),
        
        ('FR-4: Deterministic Sequential Approval Workflow Engine',
         'Memos shall follow strict linear approval stages (Step 1 -> Step 2 -> ... -> Step N). Only the currently assigned turn holder (or active delegate) can sign off. Other users cannot jump turns or approve out of order.'),
        
        ('FR-5: Mid-Stream Dynamic Reviewer Injection',
         'Active turn holders shall have the capability to use "Approve & Add Reviewer" to inject an ad-hoc specialist (e.g. Legal Counsel) between steps without breaking workflow integrity.'),
        
        ('FR-6: Change Request, Revision & Version Diffing Engine',
         'Approvers can select "Request Changes" with mandatory rationale comments. The memo status transitions to "Changes Requested" and returns to the author. When the author edits and resubmits, the system automatically creates an immutable MemoVersion snapshot.'),
        
        ('FR-7: Active Workflow Delegation Engine',
         'Users can configure temporary workflow delegations specifying a delegatee and start/end calendar dates. During active delegation, the delegatee can approve on the assignor\'s behalf with full audit attribution.'),
        
        ('FR-8: Corporate PDF Generation & Digital Seal Engine',
         'The system shall generate official memorandum PDF documents featuring company letterheads, status badges, full signature matrices, audit approval stamps, and cryptographic QR code verification seals.'),
        
        ('FR-9: In-App Real-Time Notification System',
         'The system shall dispatch in-app notifications whenever: (a) a memo requires the user\'s turn, (b) changes are requested, (c) a memo is approved or rejected, or (d) delegation authority is assigned. Users can mark alerts as read individually or in bulk.'),
        
        ('FR-10: Faceted Scoped Search & Multi-Criteria Filtering',
         'Users shall search memoranda across titles, reference numbers, author names, departments, date ranges, categories, and priority levels with real-time result highlighting.'),
        
        ('FR-11: Immutable Audit Trail & Activity Logging',
         'The system shall record an append-only audit entry for every meaningful event: user login, memo creation, step approval, rejection, rerouting, delegation, and admin configuration changes, capturing actor, timestamp, IP address, and metadata.'),
        
        ('FR-12: Executive Analytics & Turnaround Reporting',
         'Admins and executive officers shall have access to real-time analytics charts displaying status distribution, turnaround duration per department, bottleneck identification, and monthly volume throughput.')
    ]

    for fr_title, fr_desc in frs:
        h = add_section_header(fr_title, level=2)
        h.style.font.color.rgb = RGBColor(49, 46, 129)
        p = doc.add_paragraph(fr_desc)
        p.paragraph_format.line_spacing = 1.15

    doc.add_page_break()

    # -------------------------------------------------------------
    # 5. NON-FUNCTIONAL REQUIREMENTS (NFRs)
    # -------------------------------------------------------------
    add_section_header('4. Non-Functional Requirements (NFRs)', level=1)

    nfrs = [
        ('NFR-1: Performance & Low Latency', 'All REST API endpoints shall respond within 150ms under standard operational loads. Database queries must utilize indexed foreign keys and connection pooling on Neon Cloud.'),
        ('NFR-2: Strict Security & Cryptographic Integrity', 'Passwords must be salted and hashed using PBKDF2-SHA256 / Bcrypt. API sessions must use stateless JWT tokens with 24h expiration. All SQL interactions must use parameterized ORM statements to prevent SQL injection. Input HTML must be sanitized against XSS.'),
        ('NFR-3: Cross-Tenant Perimeter Enforcement', 'Any API request attempting to query or mutate an entity belonging to a different org_id must immediately fail with an HTTP 404/403 status code.'),
        ('NFR-4: High Availability & Serverless Scalability', 'The backend must operate statelessly across serverless compute instances on Vercel with automatic failover and 99.9% uptime SLA.'),
        ('NFR-5: Mobile Viewport & Touch Optimization', 'The SPA must render without horizontal scrolling or clipping across viewport widths from 320px (iPhone SE) to 4K desktop displays. Form controls and buttons must have minimum 44px touch targets.'),
        ('NFR-6: Document Reproducibility & Zero Data Loss', 'PDF generation must produce identical output deterministically. Database transactions must follow ACID principles.')
    ]

    for nfr_title, nfr_desc in nfrs:
        p = doc.add_paragraph(style='List Bullet')
        r_t = p.add_run(f'{nfr_title}: ')
        r_t.bold = True
        p.add_run(nfr_desc)

    # -------------------------------------------------------------
    # 6. SYSTEM ARCHITECTURE & DATA DESIGN
    # -------------------------------------------------------------
    add_section_header('5. System Architecture & Technical Specifications', level=1)

    doc.add_paragraph(
        '5.1 Architectural Topology:\n'
        '• Client Tier: Vanilla JavaScript Single-Page Application (SPA) with reactive routing, Tailwind CSS styling, Quill.js rich text editor, Chart.js visualizer, and Lucide icons.\n'
        '• API Application Tier: Asynchronous FastAPI (Python 3.10+) REST services structured into modular routers (auth, memos, workflow, admin, reports, delegations, notifications).\n'
        '• Business Logic Tier: Domain services for workflow state transitions, version snapshots, audit event logging, PDF rendering, and notification dispatching.\n'
        '• Data & Storage Tier: PostgreSQL 16 relational database on Neon Cloud with SSL/TLS encryption.'
    )

    doc.add_paragraph('5.2 Primary Database Schema Entities:')
    entities = [
        ('organizations', 'Master tenant record: id, name, code, contact_email, created_at.'),
        ('users', 'User profile: id, org_id, department_id, full_name, email, password_hash, role, is_active.'),
        ('departments', 'Corporate department: id, org_id, name, description.'),
        ('memo_categories', 'Taxonomy tags: id, org_id, name, description.'),
        ('memos', 'Memorandum document: id, org_id, memo_number, title, body_html, priority, status, author_id, current_assignee_id.'),
        ('memo_workflow_steps', 'Sequential stage: id, memo_id, step_number, assigned_user_id, status, decision, comments, action_timestamp.'),
        ('memo_versions', 'Immutable snapshot: id, memo_id, version_number, title, body_html, change_summary, created_by_id, created_at.'),
        ('workflow_delegations', 'Delegation rule: id, org_id, delegator_id, delegatee_id, start_date, end_date, is_active.'),
        ('audit_logs', 'Activity log: id, org_id, user_id, event_type, object_type, object_id, description, ip_address, created_at.'),
        ('notifications', 'In-app alert: id, org_id, user_id, memo_id, title, message, is_read, created_at.')
    ]

    e_table = doc.add_table(rows=len(entities)+1, cols=2)
    e_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    e_headers = ['Database Entity Table', 'Entity Description & Schema Fields']
    for j, h in enumerate(e_headers):
        cell = e_table.rows[0].cells[j]
        set_cell_background(cell, '1E1B4B')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    for i, (t_name, t_desc) in enumerate(entities):
        row = e_table.rows[i+1]
        for col_idx, text in enumerate([t_name, t_desc]):
            cell = row.cells[col_idx]
            if i % 2 == 0:
                set_cell_background(cell, 'F8FAFC')
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(8.5)

    doc.add_page_break()

    # -------------------------------------------------------------
    # 7. DEPLOYMENT & VERIFICATION MATRIX
    # -------------------------------------------------------------
    add_section_header('6. Deployment, Verification & Live Accounts', level=1)

    doc.add_paragraph('6.1 Production Cloud Deployment Details:')
    doc.add_paragraph('• Production URL: https://memo-system-pjbj.vercel.app', style='List Bullet')
    doc.add_paragraph('• Source Code ZIP: https://github.com/ZamanShafin/memo-system/raw/main/source_code.zip', style='List Bullet')
    doc.add_paragraph('• Automated Test Results: 20/20 Passing Automated Unit/Integration Tests (pytest)', style='List Bullet')

    doc.add_paragraph('6.2 Pre-Configured Test Accounts (Universal Password: password123, Org: acme):')
    
    cred_table = doc.add_table(rows=8, cols=3)
    cred_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_headers = ['Role / Persona', 'Login Email', 'Evaluation Function']
    for j, h in enumerate(c_headers):
        cell = cred_table.rows[0].cells[j]
        set_cell_background(cell, '312E81')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    creds = [
        ('Sarah Jenkins', 'admin@acmecorp.com', 'System Admin: User & Department management, audit log inspection.'),
        ('Alex Morgan', 'alex.morgan@acmecorp.com', 'Author / Requester: Submitting memos, tracking status, resubmissions.'),
        ('David Vance', 'head.eng@acmecorp.com', 'Dept Head: Tier-1 approvals, change requests, temporary delegations.'),
        ('Jessica Taylor', 'jessica.taylor@acmecorp.com', 'Acting Delegate: Approving on David Vance\'s behalf during active delegation.'),
        ('Rachel Green', 'finance.mgr@acmecorp.com', 'Finance Manager: Tier-2 financial reviews, budget vetting, rejections.'),
        ('Marcus Sterling', 'director@acmecorp.com', 'Operations Director: Tier-3 operational sign-offs, dynamic routing.'),
        ('Eleanor Vance', 'ceo@acmecorp.com', 'CEO: Final corporate approval, PDF seal generation & archiving.')
    ]

    for i, (p_name, p_email, p_func) in enumerate(creds):
        row = cred_table.rows[i+1]
        for col_idx, text in enumerate([p_name, p_email, p_func]):
            cell = row.cells[col_idx]
            if i % 2 == 0:
                set_cell_background(cell, 'F8FAFC')
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)

    doc.save('submission/PRODUCT_REQUIREMENTS_DOCUMENT.docx')
    doc.save('PRODUCT_REQUIREMENTS_DOCUMENT.docx')
    print('SUCCESS: Created PRODUCT_REQUIREMENTS_DOCUMENT.docx')

generate_prd()
