import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_submission_form_doc():
    doc = docx.Document()
    
    # Page setup - Normal margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def set_cell_background(cell, fill_color):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
        tcPr.append(shd)

    # 1. HEADER / TITLE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('CSE226 Foundations of Vibe Coding\nFinal Project Submission Form')
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 27, 75) # Indigo 950

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run('Official Course Submission Deliverable — Section 29 Compliance')
    sub_run.font.size = Pt(12)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo 600

    doc.add_paragraph()

    # METADATA TABLE
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ('Course Title:', 'CSE226 Foundations of Vibe Coding'),
        ('Institution:', 'North South University'),
        ('Student / Submitter:', 'Zaman Shafin'),
        ('Submission Date:', 'August 2026'),
        ('Project Title:', 'Multi-Tenant Inter-Office Memo Management System')
    ]
    for i, (k, v) in enumerate(meta_data):
        cell_k, cell_v = meta_table.rows[i].cells
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, 'F1F5F9')
        p_k = cell_k.paragraphs[0]
        r_k = p_k.add_run(k)
        r_k.bold = True
        r_k.font.size = Pt(10)
        p_v = cell_v.paragraphs[0]
        r_v = p_v.add_run(v)
        r_v.font.size = Pt(10)

    doc.add_paragraph()

    # SUBMISSION ITEMS (A through E)

    # ITEM A
    h_a = doc.add_heading('A. Deployed Application', level=2)
    h_a.style.font.color.rgb = RGBColor(30, 27, 75)
    p_a = doc.add_paragraph()
    r_a1 = p_a.add_run('Publicly Reachable URL: ')
    r_a1.bold = True
    r_a2 = p_a.add_run('https://memo-system-pjbj.vercel.app\n')
    r_a2.font.color.rgb = RGBColor(79, 70, 229)
    r_a2.bold = True
    p_a.add_run('Status: Live, functional, production-ready on Vercel Serverless runtime backed by Neon PostgreSQL Cloud Database.')

    # ITEM B
    h_b = doc.add_heading('B. Project Documentation', level=2)
    h_b.style.font.color.rgb = RGBColor(30, 27, 75)
    p_b = doc.add_paragraph()
    r_b1 = p_b.add_run('Documentation URL: ')
    r_b1.bold = True
    r_b2 = p_b.add_run('https://github.com/ZamanShafin/memo-system/blob/main/submission/2_PROJECT_DOCUMENTATION.md\n')
    r_b2.font.color.rgb = RGBColor(79, 70, 229)
    p_b.add_run('Description: Comprehensive report detailing system overview (26.1), requirements compliance matrix (26.2), technology stack (26.3), 3-tier architecture (26.4), database schema & multi-tenancy (26.5), sequential workflow state engine & delegation (26.6), security & RBAC (26.7), vibe-coding process (26.8), known limitations (26.9), and deployment info (26.10).')

    # ITEM C
    h_c = doc.add_heading('C. Source Code Archive', level=2)
    h_c.style.font.color.rgb = RGBColor(30, 27, 75)
    p_c = doc.add_paragraph()
    r_c1 = p_c.add_run('Source Code ZIP URL: ')
    r_c1.bold = True
    r_c2 = p_c.add_run('https://github.com/ZamanShafin/memo-system/raw/main/source_code.zip\n')
    r_c2.font.color.rgb = RGBColor(79, 70, 229)
    r_c3 = p_c.add_run('GitHub Repository: ')
    r_c3.bold = True
    r_c4 = p_c.add_run('https://github.com/ZamanShafin/memo-system\n')
    r_c4.font.color.rgb = RGBColor(79, 70, 229)
    p_c.add_run('Contents: Complete backend (FastAPI), frontend SPA (Tailwind + Vanilla JS), database schemas, seed data, configurations, dependency definitions (pyproject.toml, requirements.txt), automated test suite (pytest), and complete installation/setup instructions.')

    # ITEM D
    h_d = doc.add_heading('D. AI Prompt and Response History', level=2)
    h_d.style.font.color.rgb = RGBColor(30, 27, 75)
    p_d = doc.add_paragraph()
    r_d1 = p_d.add_run('AI Prompt/Response History URL: ')
    r_d1.bold = True
    r_d2 = p_d.add_run('https://github.com/ZamanShafin/memo-system/blob/main/submission/5_AI_PROMPT_AND_RESPONSE_LOG.md\n')
    r_d2.font.color.rgb = RGBColor(79, 70, 229)
    p_d.add_run('Description: Full chronological record of AI interactions, prompts, generated code, error debugging, UI refinements, test validations, and iterative improvements throughout the vibe-coding lifecycle with confidential credentials redacted.')

    doc.add_page_break()

    # ITEM E
    h_e = doc.add_heading('E. Demonstration Credentials', level=2)
    h_e.style.font.color.rgb = RGBColor(30, 27, 75)
    
    p_e = doc.add_paragraph()
    p_e.add_run('Organization Identifier / Code: ').bold = True
    r_org = p_e.add_run('acme\n')
    r_org.bold = True
    r_org.font.color.rgb = RGBColor(79, 70, 229)
    p_e.add_run('Universal Password for All Accounts: ').bold = True
    r_pwd = p_e.add_run('password123\n')
    r_pwd.bold = True
    r_pwd.font.color.rgb = RGBColor(225, 29, 72)
    p_e.add_run('Note on Fast Evaluation: The deployed application includes a 1-Click "Demo Switcher" dropdown in the top header, allowing evaluators to switch between personas instantly without typing credentials.')

    # CREDENTIALS TABLE
    cred_table = doc.add_table(rows=8, cols=4)
    cred_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_headers = ['Role / Persona', 'Login Email', 'Department', 'Evaluation Scenario']
    for j, h in enumerate(c_headers):
        cell = cred_table.rows[0].cells[j]
        set_cell_background(cell, '1E1B4B')
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    creds = [
        ('System Administrator', 'admin@acmecorp.com', 'Executive Operations (Admin)', 'User/dept management, audit logs, analytics'),
        ('Requester / Author', 'alex.morgan@acmecorp.com', 'Engineering & Tech (User)', 'Authoring memos, revisions, tracking status'),
        ('Department Head', 'head.eng@acmecorp.com', 'Engineering & Tech (User)', 'Tier-1 approvals, change requests, delegations'),
        ('Acting Delegate', 'jessica.taylor@acmecorp.com', 'Engineering & Tech (User)', 'Approving on behalf of David Vance (Delegation)'),
        ('Finance Manager', 'finance.mgr@acmecorp.com', 'Finance & Accounts (User)', 'Tier-2 budget approvals and rejections'),
        ('Director of Operations', 'director@acmecorp.com', 'Procurement & Ops (User)', 'Tier-3 operational sign-offs & dynamic routing'),
        ('Chief Executive Officer', 'ceo@acmecorp.com', 'Executive Office (User)', 'Final approval, PDF seal generation & closure')
    ]
    
    for i, (r_name, r_email, r_dept, r_scen) in enumerate(creds):
        row = cred_table.rows[i+1]
        for col_idx, text in enumerate([r_name, r_email, r_dept, r_scen]):
            cell = row.cells[col_idx]
            if i % 2 == 0:
                set_cell_background(cell, 'F8FAFC')
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9)

    doc.add_paragraph()

    # SUMMARY BOX
    sum_box = doc.add_paragraph()
    sum_box.paragraph_format.left_indent = Inches(0.2)
    sum_box.paragraph_format.right_indent = Inches(0.2)
    r_sb = sum_box.add_run('Declaration of Originality & Vibe-Coding Compliance:\n')
    r_sb.bold = True
    r_sb.font.size = Pt(10)
    r_sb_text = sum_box.add_run('This project was engineered following the Vibe Coding methodology using AI pair programming tools. All generated code has been validated with automated unit tests (20/20 passing), verified for strict tenant isolation, and deployed to production cloud infrastructure.')
    r_sb_text.font.size = Pt(9.5)
    r_sb_text.font.italic = True

    # Save to files
    doc.save('submission/FINAL_SUBMISSION_FORM.docx')
    doc.save('FINAL_SUBMISSION_FORM.docx')
    print('SUCCESS: Created submission/FINAL_SUBMISSION_FORM.docx and FINAL_SUBMISSION_FORM.docx')

create_submission_form_doc()
